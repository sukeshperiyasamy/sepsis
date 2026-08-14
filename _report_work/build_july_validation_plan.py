from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent))
from build_project_report import (
    ROOT, NAVY, BLUE, TEAL, INK, MUTED, LIGHT_BLUE, LIGHT_TEAL, LIGHT_GRAY,
    MID_GRAY, AMBER, RED, style_doc, set_header_footer, add_heading, add_p,
    add_bullet, add_numbered, add_callout, add_table, add_source_note,
    set_run_font, set_table_geometry, set_table_borders, shade_cell,
    add_page_break,
)
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUT = ROOT / "July_2026_Rapid_SERS_Validation_Plan_and_Report_Template.docx"


def build():
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    set_header_footer(sec)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("JULY 2026 MONTHLY PROGRESS REPORT WORK PLAN")
    set_run_font(r, name="Aptos Display", size=14, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rapid Validation of an Ag/Si SERS Substrate\nfor NAM Detection")
    set_run_font(r, name="Aptos Display", size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Five-day execution plan and honest report template")
    set_run_font(r, size=12, color=MUTED, italic=True)
    doc.add_paragraph()
    meta = [
        ("Student", "Sukesh P | M24IM1007"),
        ("Project", "Sepsis Biomarkers Analysis Using a SERS-Based Sensor"),
        ("Substrate", "Silver-coated nanostructured silicon (Ag/Si)"),
        ("Reporting month", "July 2026"),
        ("Status", "To be completed using actual measurements before submission"),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2500, 6860], indent_dxa=0)
    set_table_borders(table, color=MID_GRAY, size=5)
    for label, value in meta:
        row = table.add_row()
        shade_cell(row.cells[0], LIGHT_GRAY)
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.style = doc.styles["Table Text"]
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=NAVY if i == 0 else INK, bold=(i == 0))
    set_table_geometry(table, [2500, 6860], indent_dxa=0)
    add_callout(doc, "Integrity requirement", "This document is a work plan and report template. Do not write an experiment as completed until it has been performed and the actual raw data, processed data, plots, and values have been saved. Use phrases such as 'preliminary validation', 'observed under the tested condition', and 'requires further optimization' where appropriate.", fill=AMBER, accent="9A6700")

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Cm(21.0)
    body.page_height = Cm(29.7)
    body.top_margin = Cm(2.0)
    body.bottom_margin = Cm(1.8)
    body.left_margin = Cm(1.9)
    body.right_margin = Cm(1.9)
    body.header_distance = Cm(1.0)
    body.footer_distance = Cm(0.9)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    set_header_footer(body)

    add_heading(doc, "1. July Objective", 1)
    add_p(doc, "The July objective is to perform a rapid, controlled validation study of the existing silver-coated nanostructured silicon SERS substrate. R6G will be used as a positive control to verify substrate performance, followed by preliminary NAM measurements to determine whether the substrate produces a reproducible and concentration-dependent bacterial biomarker response.")
    add_callout(doc, "Scope control", "For the five-day period, focus on one substrate, one control molecule, and one biomarker. The priority is a defensible NAM validation result, not a full clinical diagnostic study.", fill=LIGHT_TEAL, accent=TEAL)

    add_heading(doc, "2. Experiments to Perform", 1)
    add_table(doc, ["Experiment", "Purpose", "Minimum design", "Primary outputs"], [
        ("E1. R6G substrate QC", "Confirm that the current Ag/Si substrate is Raman/SERS-active before interpreting NAM.", "Blank silicon/blank Ag-Si if available + R6G on at least 3 spots; use current starting condition: 785 nm, 10% power, 5 accumulations, 6-minute soaking.", "R6G spectrum, SNR, peak positions, FWHM, peak matching, spot-to-spot CV/correlation."),
        ("E2. NAM blank and reference", "Separate NAM response from substrate background and solvent background.", "Ag/Si blank, solvent blank, and NAM at one approved reference concentration; at least 3 spots per condition.", "Blank-subtracted spectrum, NAM peak list, SNR above blank, repeatability."),
        ("E3. NAM concentration response", "Test whether NAM signal changes systematically with concentration.", "Blank + low, medium, and high approved NAM concentrations; at least 3 spots per level; keep acquisition settings fixed.", "Peak area/height versus concentration, SNR versus concentration, preliminary response trend."),
        ("E4. NAM deposition/adsorption comparison", "Identify whether weak signal is caused by poor analyte-surface interaction.", "Use the same NAM concentration and compare 2-3 approved deposition or adsorption/drying conditions; at least 3 spots per condition.", "Signal intensity, peak presence, SNR, and reproducibility by preparation condition."),
        ("E5. Acquisition check", "Check whether integration time or accumulations improve NAM without excessive background or damage.", "Compare two instrument settings only, keeping sample and substrate fixed; include a blank and repeated NAM spot.", "SNR, baseline stability, peak sharpness, acquisition time, and visual sample condition."),
        ("E6. Repeatability check", "Estimate short-term measurement and substrate variability.", "Repeat the best condition on 3 spots and, if available, 2-3 substrate pieces or batches.", "Mean, standard deviation, CV, Pearson correlation, and pass/fail decision."),
    ], [1450, 2450, 3100, 2360], font_size=8.2)
    add_callout(doc, "Priority order", "If instrument time is limited, complete E1, E2, E3, and E6 first. E4 and E5 are secondary optimization experiments and should not replace the blank, concentration, and repeatability measurements.", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "3. Five-Day Execution Schedule", 1)
    add_table(doc, ["Day", "Work to complete", "End-of-day evidence"], [
        ("Day 1", "Prepare/confirm approved stocks, label substrate pieces, record instrument settings, measure blanks and R6G controls.", "Sample log, substrate IDs, raw R6G and blank files, first QC plot."),
        ("Day 2", "Measure NAM reference condition and concentration series using fixed acquisition settings.", "Raw NAM files for blank/low/medium/high levels and acquisition log."),
        ("Day 3", "Repeat the strongest NAM conditions; compare deposition/adsorption or drying conditions if time permits.", "Replicate spectra, preparation comparison table, preliminary SNR results."),
        ("Day 4", "Complete short-term repeatability and acquisition check; process all data using one locked pipeline.", "Processed CSVs, peak table, SNR/CV/correlation table, overlay figures."),
        ("Day 5", "Interpret results, document limitations, prepare figures and monthly report, and obtain supervisor review.", "Final report draft, appendix of raw-file names, conclusion, and next-step decision."),
    ], [1200, 5700, 2460], font_size=8.8)

    add_heading(doc, "4. Comparison and Validation Strategy", 1)
    add_heading(doc, "4.1 Required comparisons", 2)
    add_table(doc, ["Comparison", "Why it is needed", "How to interpret"], [
        ("R6G on Ag/Si versus blank substrate", "Shows whether the substrate provides enhancement above background.", "R6G peaks and higher SNR support substrate functionality."),
        ("NAM on Ag/Si versus Ag/Si blank", "Shows whether observed peaks are associated with NAM rather than the substrate.", "NAM features must exceed the blank and recur across spots."),
        ("NAM concentrations versus blank", "Tests concentration dependence.", "A monotonic or statistically supported trend supports further LOD work."),
        ("Different deposition/adsorption conditions", "Tests whether surface contact controls the weak NAM response.", "The best condition is the one with the strongest repeatable blank-subtracted signal."),
        ("Replicate spots/batches", "Tests reproducibility and substrate uniformity.", "Low variation and high correlation support robustness; high variation indicates substrate or preparation heterogeneity."),
        ("Two acquisition settings", "Separates instrument-limited signal from sample/interface-limited signal.", "Improved SNR with stable baseline supports an acquisition adjustment; no improvement points to sample/substrate optimization."),
    ], [2500, 3600, 3260], font_size=8.7)
    add_heading(doc, "4.2 Metrics to calculate", 2)
    for item in [
        "SNR relative to the same blank definition for all conditions.",
        "Peak position and peak intensity/area for the selected NAM fingerprint regions.",
        "FWHM or peak sharpness where the peak is sufficiently defined.",
        "Mean, standard deviation, and coefficient of variation across replicate spots.",
        "Pearson correlation between normalized replicate spectra, with the comparison range stated.",
        "Blank-subtracted peak area or height versus NAM concentration.",
        "Preliminary detection threshold defined from the blank distribution; do not claim a formal LOD unless the required replicate and statistical design has been completed.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Validation language", "A five-day study can support preliminary validation of substrate function, signal detectability, short-term repeatability, and concentration response. It cannot by itself establish clinical sensitivity, specificity, clinical LOD, or diagnostic accuracy.", fill=RED, accent="A33A3A")

    add_heading(doc, "5. Data Processing Workflow", 1)
    for item in [
        "Create a sample manifest before measurement with substrate ID, sample ID, concentration, deposition condition, spot number, laser power, integration time, accumulations, and file name.",
        "Keep raw files unchanged and save processed outputs in a separate folder.",
        "Apply the same column parsing, Raman range, baseline correction, smoothing, and normalization to every condition.",
        "Use the same blank definition for SNR and blank subtraction throughout the comparison.",
        "Generate overlays for blank versus R6G and blank versus NAM, concentration-response plots, replicate correlation/variation plots, and a summary table.",
        "Record failed measurements and exclude them only with a written reason; do not silently remove unfavorable spectra.",
    ]:
        add_numbered(doc, item)
    add_table(doc, ["File to save", "Required contents"], [
        ("raw_data_manifest.csv", "Sample ID, substrate ID, condition, spot, acquisition settings, date/time, raw filename."),
        ("processed_spectra.csv", "Raman shift, raw intensity, corrected intensity, smoothed intensity, normalized intensity."),
        ("peak_summary.csv", "Condition, spot, peak position, peak intensity/area, SNR, FWHM, notes."),
        ("replicate_statistics.csv", "Mean, SD, CV, correlation, number of spots, pass/fail status."),
        ("figures/", "Blank/R6G overlay, blank/NAM overlay, concentration trend, replicate overlay, and representative spectra."),
        ("experiment_log.docx or .xlsx", "Sample preparation, deviations, instrument issues, and observations."),
    ], [2600, 6760], font_size=8.8)

    add_heading(doc, "6. Expected Outcomes and Decision Table", 1)
    add_table(doc, ["Observed result after the experiment", "Conclusion to write", "Next action"], [
        ("R6G is strong and reproducible; NAM is above blank at several spots.", "The Ag/Si substrate supports preliminary NAM detection under the tested condition.", "Repeat the best condition, then begin a formal concentration/LOD study."),
        ("R6G is strong; NAM is weak but improves with concentration or deposition condition.", "NAM response is preliminary and sample-interface optimization is still required.", "Optimize the best preparation variable and repeat with more replicates."),
        ("R6G is strong; NAM is visible but highly variable between spots.", "The main limitation appears to be spot-to-spot or substrate heterogeneity.", "Improve substrate uniformity, use more spots/batches, and quantify CV/correlation."),
        ("R6G is weak or inconsistent.", "The current substrate/acquisition system did not pass the control check.", "Do not interpret NAM; troubleshoot substrate preparation and instrument alignment first."),
        ("Neither R6G nor NAM is reliable.", "The measurement run was inconclusive and cannot support biomarker validation.", "Document the failure, review safety and instrument logs, and repeat after supervisor review."),
    ], [3500, 3900, 1960], font_size=8.6)
    add_callout(doc, "Do not force a positive conclusion", "A result showing that NAM remains weak is still useful. It validates the current limitation, identifies the bottleneck, and justifies the next optimization step. Report the actual outcome with the measured numbers.", fill=AMBER, accent="9A6700")

    add_heading(doc, "7. July Report Writing Template", 1)
    add_heading(doc, "7.1 Work completed during July", 2)
    add_p(doc, "During July 2026, a preliminary validation study was carried out to evaluate the performance of a silver-coated nanostructured silicon SERS substrate for bacterial biomarker analysis. Rhodamine 6G was used as a positive-control molecule, followed by measurements of N-acetylmuramic acid (NAM) under controlled sample and acquisition conditions. The study compared blank substrate response, control-molecule enhancement, NAM signal, concentration dependence, and short-term repeatability.")
    add_p(doc, "Replace the preceding paragraph only after the corresponding experiments have actually been performed. Add the number of substrate pieces, spots, concentrations, acquisition settings, and dates from the experiment log.", italic=True, color=MUTED)
    add_heading(doc, "7.2 Results table to fill with actual values", 2)
    add_table(doc, ["Condition", "n", "SNR mean +/- SD", "Peak/area result", "CV/correlation", "Interpretation"], [
        ("Blank silicon", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("Blank Ag/Si", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("R6G control", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("NAM blank/solvent", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("NAM low concentration", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("NAM medium concentration", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("NAM high concentration", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("Best repeatability condition", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"),
    ], [1700, 600, 1600, 2100, 1700, 1660], font_size=8.2)
    add_heading(doc, "7.3 Honest conclusion templates", 2)
    add_bullet(doc, "Positive preliminary result: 'Under the tested conditions, the Ag/Si substrate produced a measurable and repeatable NAM-associated response above the blank. The result is preliminary and requires concentration-series expansion and independent batch validation.'")
    add_bullet(doc, "Optimization result: 'The control experiment confirmed substrate activity, while NAM response remained weak/variable. The data indicate that analyte adsorption, deposition, or substrate uniformity remains the limiting factor.'")
    add_bullet(doc, "Inconclusive result: 'The validation run did not provide sufficient evidence for a reproducible NAM response. The control and measurement logs identify the next troubleshooting steps, and the study will be repeated after method review.'")
    add_p(doc, "Choose only the conclusion supported by the actual data. Do not copy a positive template if the measurements do not support it.", italic=True, color=MUTED)

    add_heading(doc, "8. Safety and Recordkeeping", 1)
    for item in [
        "Use only approved laboratory SOPs and supervisor-authorized conditions for HF etching, silver deposition, Raman operation, and waste disposal.",
        "If existing Ag/Si substrates are available, prioritize measurement and validation rather than starting an unapproved fabrication change during the five-day period.",
        "Record all deviations, failed spots, instrument warnings, sample damage, and discarded measurements.",
        "Keep raw data read-only and preserve the original file names so the July report can be audited.",
        "Do not describe a planned experiment as completed; use 'planned', 'in progress', or 'not completed' where appropriate.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Immediate Priority", 1)
    add_p(doc, "The highest-value experiment for this reporting period is a controlled NAM validation on the existing Ag/Si substrate with R6G quality control, blanks, a small concentration series, and replicate spots. The report should show what was tested, what was measured, what failed or succeeded, and what the result implies for the next phase.")
    add_callout(doc, "Five-day target", "Produce one complete, traceable validation package: raw data manifest + processed spectra + quantitative comparison table + four figures + conclusion with limitations + supervisor-reviewed next step.", fill=LIGHT_TEAL, accent=TEAL)

    props = doc.core_properties
    props.title = "July 2026 Rapid SERS Validation Plan and Monthly Report Template"
    props.subject = "Five-day Ag/Si substrate and NAM validation plan"
    props.author = "Sukesh P"
    props.keywords = "July 2026, SERS, Ag/Si, NAM, validation, monthly progress report"
    props.comments = "Template and execution plan. Complete with actual measured data before submission."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
