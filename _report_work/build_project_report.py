from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\sukes\Downloads\mtp")
OUT = ROOT / "Project_Status_Report_Sepsis_SERS_As_of_August_2026.docx"

NAVY = "123B5D"
BLUE = "2E74B5"
TEAL = "167D8D"
INK = "1F2933"
MUTED = "5B6770"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E9F5F5"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9E0E6"
WHITE = "FFFFFF"
AMBER = "FFF4D6"
RED = "FBEAEA"
GREEN = "EAF6EE"


def set_run_font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    ppr.append(node)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_bottom_border(paragraph, color=BLUE, size=10, space=4):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(0.9)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, TEAL, 9, 4),
    ):
        st = doc.styles[name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Table Text" not in doc.styles:
        st = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = doc.styles["Table Text"]
    st.font.name = "Aptos"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    st.font.size = Pt(9.0)
    st.font.color.rgb = RGBColor.from_string(INK)
    st.paragraph_format.space_after = Pt(2)
    st.paragraph_format.line_spacing = 1.0

    if "Caption Custom" not in doc.styles:
        st = doc.styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = doc.styles["Caption Custom"]
    st.font.name = "Aptos"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    st.font.size = Pt(8.5)
    st.font.italic = True
    st.font.color.rgb = RGBColor.from_string(MUTED)
    st.paragraph_format.space_before = Pt(3)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SEPSIS BIOMARKERS ANALYSIS  |  SERS PROJECT STATUS REPORT")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    add_bottom_border(p, color=MID_GRAY, size=5, space=2)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("As of 03 August 2026  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(p)


def add_p(doc, text="", *, bold_prefix=None, italic=False, align=None, after=6, before=0, color=INK):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.32 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=0)
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_run_font(r, size=9.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int], *, header_fill=LIGHT_BLUE, font_size=8.8):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent_dxa=0)
    set_table_borders(table, color=MID_GRAY, size=5)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row_vals in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_vals):
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths, indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, image_path: Path, caption: str, width=6.25):
    if not image_path.exists():
        add_callout(doc, "Figure unavailable", f"Expected source figure was not found: {image_path}", fill=AMBER, accent="9A6700")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width))
    # Add descriptive alt text so the report remains accessible when figures are read by assistive technology.
    descr = caption.split(" Source:", 1)[0]
    inline._inline.docPr.set("descr", descr)
    inline._inline.docPr.set("title", descr)
    cap = doc.add_paragraph(caption, style="Caption Custom")
    return cap


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Evidence source: ")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=8.5, color=MUTED, italic=True)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_report():
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    # Use an academic A4 submission geometry as a named project-context override.
    set_header_footer(sec)

    # Cover page
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("PROJECT STATUS REPORT")
    set_run_font(r, name="Aptos Display", size=14, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Sepsis Biomarkers Analysis Using a\nSERS-Based Sensor")
    set_run_font(r, name="Aptos Display", size=26, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Consolidated report of work completed to date")
    set_run_font(r, size=13, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    add_bottom_border(p, color=TEAL, size=14, space=1)

    meta = [
        ("Prepared by", "Sukesh P"),
        ("Programme / ID", "Masters in Medical Technologies | M24IM1007"),
        ("Institution", "IIT Jodhpur and AIIMS Jodhpur Joint Programme"),
        ("Evidence cutoff", "03 August 2026"),
        ("Report purpose", "Submission-ready record of completed work, results, limitations, and next steps"),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2400, 6960], indent_dxa=0)
    set_table_borders(table, color=MID_GRAY, size=5)
    for label, value in meta:
        row = table.add_row()
        shade_cell(row.cells[0], LIGHT_GRAY)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            p.style = doc.styles["Table Text"]
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=NAVY if idx == 0 else INK, bold=(idx == 0))
    set_table_geometry(table, [2400, 6960], indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Compiled from the project folder contents")
    set_run_font(r, size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This report distinguishes completed evidence, preliminary results, and proposed future work.")
    set_run_font(r, size=9, color=MUTED, italic=True)

    # Body section
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    body_sec.page_width = Cm(21.0)
    body_sec.page_height = Cm(29.7)
    body_sec.top_margin = Cm(2.0)
    body_sec.bottom_margin = Cm(1.8)
    body_sec.left_margin = Cm(1.9)
    body_sec.right_margin = Cm(1.9)
    body_sec.header_distance = Cm(1.0)
    body_sec.footer_distance = Cm(0.9)
    body_sec.header.is_linked_to_previous = False
    body_sec.footer.is_linked_to_previous = False
    set_header_footer(body_sec)

    add_heading(doc, "Executive Summary", 1)
    add_p(doc, "The project has established a combined experimental, computational, and data-analysis foundation for a Surface-Enhanced Raman Spectroscopy (SERS) platform aimed at rapid, label-free analysis of bacterial sepsis biomarkers. The work has moved beyond literature review and conceptual design into reference-spectrum generation, controlled Raman data acquisition, DFT-based molecular modelling, automated spectral preprocessing, fragment-based reconstruction of complex biomarkers, and initial SERS-substrate validation.")
    add_p(doc, "Three bacterial cell-wall-associated targets were investigated: N-acetylmuramic acid (NAM) as a peptidoglycan-related marker, lipoteichoic acid (LTA) as a Gram-positive cell-wall marker, and lipopolysaccharide (LPS) as a Gram-negative outer-membrane marker. The project also developed a silver-coated nanostructured silicon substrate, validated first with Rhodamine 6G (R6G), and then tested preliminarily with NAM.")
    add_callout(doc, "Current status", "The project is scientifically established at the proof-of-concept and reference-data stage. The main remaining work is to optimize NAM enhancement, complete concentration-dependent and replicate testing, reconcile the different analysis branches, and validate the platform using real clinical samples.", fill=LIGHT_TEAL, accent=TEAL)
    add_p(doc, "The strongest completed evidence is the NAM peak-position validation and the reproducible experimental LTA dataset. The LPS workflow is complete as an analytical pipeline, but the current fit is modest and should be treated as an interpretable preliminary model. The silver-coated silicon substrate is functional for R6G, while NAM remains detectable but insufficiently enhanced for a final diagnostic claim.")

    add_heading(doc, "Report Contents", 1)
    contents = [
        "1. Scope, evidence basis, and folder inventory",
        "2. Project aim and integrated technical framework",
        "3. Chronological progress to date",
        "4. Detailed technical work completed",
        "5. Results by biomarker and sensor component",
        "6. Deliverables produced",
        "7. Current status, limitations, and evidence gaps",
        "8. Recommended next steps",
        "9. Overall conclusion",
        "Appendix A. File inventory and evidence map",
        "Appendix B. Key quantitative results",
    ]
    for item in contents:
        add_bullet(doc, item)

    add_heading(doc, "1. Scope, Evidence Basis, and Folder Inventory", 1)
    add_p(doc, "This report was prepared by reviewing the project folder as a whole, with emphasis on project-authored reports, notebooks, scripts, raw and processed Raman data, output tables, figures, presentations, and supporting protocol documents. Reference papers, administrative forms, and bundled Python environments were treated as supporting context rather than as experimental results.")
    add_table(doc, ["Inventory category", "Observed contents", "Interpretation for this report"], [
        ("Project folder", "1,924 files excluding packaged environment caches", "Large working archive containing raw data, analysis branches, reports, presentations, and generated outputs."),
        ("Experimental and processed data", "920 CSV files and 462 XLSX files", "Multiple acquisition campaigns for NAM, LTA, LPS, and parameter optimization."),
        ("Documentation and presentation", "52 DOCX, 78 PDF, 15 PPTX, 43 TXT, 4 MD", "Progress reports, technical reports, review material, SOPs, proposals, and literature notes."),
        ("Analysis code", "40 PY and 24 IPYNB files", "Data cleaning, spectral preprocessing, plotting, DFT post-processing, peak assignment, validation, and fitting."),
        ("Figures and vector outputs", "195 PNG, 25 SVG, plus JPG/GIF/WEBP outputs", "Publication-style overlays, heatmaps, spectral comparisons, and review figures."),
    ], [1900, 2450, 5010])
    add_callout(doc, "Interpretation note", "The folder contains several iterative versions of reports and analysis branches. Where values differ, this report identifies the dataset or output file associated with the number and preserves the distinction between a completed workflow and a final validated claim.", fill=AMBER, accent="9A6700")
    add_source_note(doc, "Folder inventory generated from the project directory; primary evidence includes 20-04, review meeting, molecule building, python code, report final masters, repot, and top-level project reports.")

    add_heading(doc, "2. Project Aim and Integrated Technical Framework", 1)
    add_p(doc, "The central aim is to develop a rapid SERS-based sensing approach for bacterial sepsis-associated molecular signatures. The project responds to the clinical limitations of blood culture, immunoassays, and molecular testing: they can be slow, resource-intensive, or difficult to deploy at the point of care. The proposed platform combines a plasmonic substrate, Raman acquisition, computational reference spectra, and automated analysis.")
    add_heading(doc, "2.1 Target biomarker panel", 2)
    add_table(doc, ["Target", "Biological relevance", "Modelling / measurement strategy"], [
        ("NAM", "Peptidoglycan-related bacterial cell-wall marker; used as the first detailed reference target.", "Whole-molecule DFT reference plus controlled pure-powder Raman acquisition and peak validation."),
        ("LTA", "Gram-positive bacterial cell-wall polymer containing phosphate-rich backbone and substitutions.", "Fragment-based DFT using glycerol-phosphate, D-alanine, and GlcNAc; powder Raman and reproducibility analysis."),
        ("LPS", "Gram-negative outer-membrane glycolipid associated with endotoxin-mediated inflammation.", "Fragment-based DFT using KDO, heptose, glucosamine, myristic acid, and phosphoric acid; NNLS spectral fitting."),
        ("R6G control", "Standard Raman probe used to validate substrate enhancement and wavelength calibration.", "Silver-coated nanostructured silicon substrate testing with peak matching, SNR, reproducibility, and optimization."),
    ], [1300, 3500, 4560])
    add_heading(doc, "2.2 Integrated workflow", 2)
    add_p(doc, "The documented project workflow is modular and supports both reference-spectrum development and future clinical translation:")
    for item in [
        "Problem identification and literature review on sepsis, Raman/SERS, bacterial markers, and existing diagnostic solutions.",
        "Selection of NAM, LTA, and LPS as chemically meaningful bacterial cell-wall targets.",
        "Substrate design and fabrication using nanostructured silicon with silver deposition, followed by R6G validation.",
        "Sample preparation and Raman acquisition across laser power, integration time, accumulations, and spatial locations.",
        "Data cleaning, dark subtraction, ALS baseline correction, Savitzky-Golay smoothing, normalization, peak detection, and quality scoring.",
        "DFT geometry optimization, frequency calculation, Raman activity extraction, frequency scaling, normalization, and Gaussian broadening.",
        "Peak matching, functional-group assignment, spectral overlays, similarity metrics, residual analysis, and NNLS fitting.",
        "Future integration with concentration response, clinical samples, multiplex classification, and portable point-of-care readout.",
    ]:
        add_numbered(doc, item)
    add_figure(doc, ROOT / "review meeting" / "NAM-EXP-P" / "NAM" / "output_figures" / "fig09_spectral_similarity_overlay.png", "Figure 1. Example of the integrated validation view: normalized experimental NAM spectrum, broadened DFT simulation, and residual trace. Source: review meeting/NAM-EXP-P/NAM/output_figures/fig09_spectral_similarity_overlay.png.")

    add_heading(doc, "3. Chronological Progress to Date", 1)
    add_table(doc, ["Period / phase", "Work completed", "Evidence in the folder"], [
        ("Aug-Nov 2025: foundation", "Sepsis problem definition, SERS literature review, biomarker exploration, early project reports, presentations, and initial workflow framing.", "Top-level reports, SERS presentations, sepsis background documents, literature PDFs, and early project report files."),
        ("Nov 2025-Feb 2026: NAM baseline", "Pure NAM powder measurements, acquisition-parameter exploration, Gaussian DFT reference generation, simulation-vs-experiment comparison, and early reporting.", "january report.docx, NAM documentation, NAM simulation files, NAM review documents, and processed spectra."),
        ("Mar-Apr 2026: automated biomarker analysis", "Expanded NAM ranking/validation, LTA powder reproducibility analysis, LTA fragment modelling, LPS fragment modelling, NNLS fitting, and publication-style figures/tables.", "review meeting analysis folders, notebooks, output reports, CSV/XLSX tables, and plotted figures."),
        ("Apr 2026: reporting and translation", "Consolidated project reports, review presentations, sample-biobank proposal, sample collection/storage SOP, and blood sample preparation planning.", "20-04 reports, review meeting/report doc, SOPs, proposals, and presentations."),
        ("Jun 2026: substrate validation", "Ag/Si substrate fabrication and R6G validation, preprocessing pipeline, reproducibility analysis, and preliminary NAM SERS testing.", "Development and Validation of a Silver.docx and Monthly Progress Report-June.docx."),
        ("As of Aug 2026: current status", "A substantial reference-data and analysis foundation is complete; final diagnostic validation remains pending.", "This consolidated report, based on the full folder evidence set and latest dated outputs."),
    ], [1900, 4250, 4210])

    add_heading(doc, "4. Detailed Technical Work Completed", 1)
    add_heading(doc, "4.1 Experimental data acquisition and organization", 2)
    add_p(doc, "The folder contains multiple acquisition campaigns with Raman data stored in CSV and Excel formats. File naming conventions encode integration time, laser power, replicate or accumulation information, and sample/spot identity. Several datasets contain 2048 Raman points per spectrum, with metadata rows storing integration time, excitation wavelength, and laser power.")
    add_p(doc, "The broad parameter studies examined laser power and integration time systematically. One documented analysis branch contains 80 unique spectra across 16 power levels and 5 integration times. The NAM validation branch processed 81 Excel spectra in the 200-3000 cm-1 range. The LTA detailed analysis processed 136 files across 10-25 s integration times and 5-80% power, while an earlier fixed-condition LTA branch analyzed 12 spatial spots at 20 s and 20 power.")
    add_heading(doc, "4.2 Spectral preprocessing and quality control", 2)
    for item in [
        "Metadata parsing and numeric-column detection for CSV and Excel layouts.",
        "Removal of blank rows, non-numeric entries, and invalid records; sorting by Raman shift and handling of duplicate x-values where required.",
        "Use of dark-subtracted intensity as the primary signal channel in the parameter-study data.",
        "Asymmetric Least Squares (ALS) baseline correction for fluorescence/background drift.",
        "Savitzky-Golay smoothing to reduce high-frequency noise while preserving peak shape.",
        "Cosmic-ray/spike detection and correction in the substrate-validation pipeline.",
        "Vector, area, min-max, or activity normalization depending on the analysis objective.",
        "Automated peak detection using prominence, width, spacing, and local-maximum logic.",
        "Quality scoring using SNR, peak intensity, FWHM, sharpness, baseline quality, and composite scores.",
        "Export of processed spectra, ranked tables, peak-matching tables, model metrics, residuals, and publication-ready figures.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Parameter-study result", "The enhanced Raman analysis branch reports an average power-response R2 of 0.9916, all curves above 0.97, no saturation up to 80 W and 25 s, a baseline noise estimate of approximately 114 a.u., and a best-efficiency condition of 5 W and 10 s. These results support the quality-control and acquisition-optimization part of the project, but the source notes also state that no replicate measurements were available for that branch.", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "4.3 DFT and computational modelling", 2)
    add_p(doc, "DFT modelling was used to generate molecular reference spectra and to support interpretation of experimental peaks. The workflow generally comprised molecular structure preparation, geometry optimization, frequency calculation with Raman activity, extraction of positive vibrational frequencies, scaling to reduce harmonic overestimation, normalization, and Gaussian or Lorentzian broadening before comparison with measured spectra.")
    add_table(doc, ["Modelling stream", "Completed implementation", "Purpose"], [
        ("NAM whole-molecule reference", "Gaussian-based geometry optimization and Raman frequency/activity calculation; later analysis uses a 0.96 scaling factor and broadened spectra.", "Validate the pure NAM Raman fingerprint and support peak assignment."),
        ("LTA fragment model", "Glycerol-phosphate, D-alanine, and N-acetyl-D-glucosamine simulated with B3LYP/6-31G(d), implicit water/SMD, and scaled/normalized spectra.", "Represent phosphate backbone, amino substitution, and carbohydrate contributions without simulating intact polymer LTA."),
        ("LPS fragment model", "KDO, heptose, glucosamine, myristic acid, and phosphoric acid simulated and processed into a common grid for composite fitting.", "Approximate the Raman response of a large, heterogeneous LPS molecule using interpretable structural fragments."),
        ("Substrate/control validation", "R6G theoretical/experimental comparisons and Raman peak matching were used to check calibration and enhancement performance.", "Establish that the fabricated Ag/Si substrate and measurement workflow are functional before weak biomarker testing."),
    ], [1900, 4630, 4030])
    add_callout(doc, "Method harmonization needed", "The folder contains multiple NAM simulation generations: an earlier January report records Gaussian 09 with B3LYP/6-311+G(d,p), while later NAM documentation records Gaussian 16 and a different basis set. This reflects development over time, but the final thesis should select one authoritative computational protocol and state it consistently.", fill=AMBER, accent="9A6700")

    add_heading(doc, "5. Results by Biomarker and Sensor Component", 1)
    add_heading(doc, "5.1 NAM: reference fingerprint and DFT validation", 2)
    add_p(doc, "NAM was the first target taken through a detailed combined computational and experimental workflow. Pure NAM powder was measured under a matrix of integration times and laser powers, and a separate notebook pipeline ranked spectra, selected a representative spectrum, matched experimental peaks to DFT peaks, and generated annotated outputs.")
    add_table(doc, ["Metric", "Completed result", "Meaning"], [
        ("Spectra processed", "81 experimental Excel spectra in the NAM validation branch", "Sufficient parameter coverage for quality ranking in the current dataset."),
        ("Best spectrum", "sec-25_power-80_i-2.xlsx", "25 s integration and 80 power selected by composite score."),
        ("Best SNR", "95.91", "Strongest spectral quality in the ranking branch."),
        ("Peak matching", "51 matched peaks", "Broad peak-position correspondence between experimental and DFT data."),
        ("Frequency error", "Mean +0.378 cm-1; standard deviation 2.102 cm-1; RMSE 2.115 cm-1; maximum absolute error 3.8 cm-1", "Excellent peak-position agreement under the selected 0.96 scaling factor."),
        ("Full-spectrum similarity", "Cosine similarity 0.5868; Pearson correlation 0.1710", "Poor global-shape similarity in the validation branch; peak alignment is stronger than intensity-profile agreement."),
        ("Output products", "Ranked table, processed CSVs, peak-matching/assignment tables, 13+ figures, and text report", "Reproducible analysis package for the NAM reference study."),
    ], [2200, 3600, 4760])
    add_p(doc, "The NAM work supports a scientifically useful reference spectrum, but it does not yet prove clinical specificity or final SERS performance. The folder's own conclusions correctly emphasize that NAM should be identified from a multi-peak fingerprint rather than a single peak, and that pure NAM powder is a reference control rather than a substitute for complex bacterial or clinical samples.")
    add_figure(doc, ROOT / "review meeting" / "NAM-EXP-P" / "NAM" / "output_figures" / "fig10_validation_summary_4panel.png", "Figure 2. NAM validation summary output showing peak-error distribution and experimental/DFT comparison. Source: review meeting/NAM-EXP-P/NAM/output_figures/fig10_validation_summary_4panel.png.")
    add_source_note(doc, "Primary evidence: review meeting/NAM-EXP-P/NAM/NAM_Raman_Analysis.ipynb, NAM_Validation_Analysis.ipynb, NAM_Peak_Assignment.ipynb, output_figures/NAM_analysis_report.txt, and 20-04/finalmtp.docx.")

    add_heading(doc, "5.2 LTA: experimental reproducibility and fragment interpretation", 2)
    add_p(doc, "LTA was investigated through two complementary branches. The fixed-condition powder branch analyzed 12 spatial spots at 20 s integration and 20 power, while the expanded parameter branch ranked 136 spectra across multiple acquisition settings. Fragment-based DFT modelling represented the structural domains of LTA using glycerol-phosphate, D-alanine, and GlcNAc.")
    add_table(doc, ["Metric", "Completed result", "Interpretation"], [
        ("Fixed-condition spot study", "12 spots; 7 characteristic peaks; average spectral correlation 0.9787; average peak CV 16.74%", "Good spot-to-spot reproducibility and acceptable powder homogeneity in that dataset."),
        ("Expanded acquisition study", "136 files; best file sec-20_power-40_i-2_2.xlsx", "Parameter ranking identified a high-quality representative spectrum."),
        ("Best spectrum quality", "20 s, 40 power; SNR 139.03; FWHM 9.13 cm-1; quality score 69.7490", "Strong spectral quality under the expanded study settings."),
        ("Major experimental bands", "Approx. 413, 619, 795, 1000, 1030, 1182, and 1601 cm-1", "Bands assigned to phosphate/backbone, carbohydrate, CH2, and amide-related regions."),
        ("Peak matching", "7/7 selected peaks; mean error 0.41 cm-1; RMSE 2.70 cm-1; maximum error 4.00 cm-1", "Good selected-peak alignment with the DFT reference."),
        ("Global similarity", "Cosine 0.2397; Pearson -0.0084 in the expanded report", "Whole-spectrum shape agreement remains weak despite selected peak matches."),
        ("Unresolved analysis branch", "20-04/lta/results/final_summary_report.txt contains NaN weights and NaN fit metrics", "This branch should not be used as final evidence until rerun and verified."),
    ], [2200, 3600, 4760])
    add_p(doc, "The combined evidence shows that LTA has a reproducible experimental fingerprint and a chemically interpretable fragment library. The current record is not yet a fully validated quantitative model because selected-peak agreement and global spectral similarity tell different stories. The next iteration should use a single curated dataset, consistent normalization, and a verified NNLS branch with complete non-NaN metrics.")
    add_figure(doc, ROOT / "20-04" / "20sec20power-variable location LTA" / "Spectral_Correlation_Matrix.png", "Figure 3. LTA spot-to-spot Pearson correlation matrix from the fixed-condition 12-spot analysis. Source: 20-04/20sec20power-variable location LTA/Spectral_Correlation_Matrix.png.")
    add_figure(doc, ROOT / "20-04" / "lta" / "results" / "final_overlay.png", "Figure 4. LTA experimental spectrum compared with manual and NNLS fragment models. The plotted fit metrics illustrate why the global model requires further refinement. Source: 20-04/lta/results/final_overlay.png.")

    add_heading(doc, "5.3 LPS: fragment-based reconstruction and NNLS fitting", 2)
    add_p(doc, "LPS was treated as a large and heterogeneous glycolipid for which intact-molecule quantum simulation is computationally impractical. Five fragments were selected to represent core sugar, amino-sugar, lipid, and phosphate contributions. The final documented workflow used an experimental spectrum named sec-60_power-20_i-30 and applied scaling, ALS correction, Savitzky-Golay smoothing, min-max normalization, common-grid interpolation, adaptive peak detection, and non-negative least squares fitting.")
    add_table(doc, ["Metric", "Completed result", "Interpretation"], [
        ("DFT fragments", "KDO, heptose, glucosamine, myristic acid, phosphoric acid", "Interpretable structural basis for a composite LPS spectrum."),
        ("Analysis range", "200-2000 cm-1 at 1 cm-1 spacing; 1,801 grid points", "Common comparison grid for fitting and residual analysis."),
        ("NNLS weights", "KDO 47.3%; heptose 25.8%; glucosamine 15.3%; myristic acid 11.7%; phosphoric acid 0.0%", "KDO is the dominant fitted contributor in the current model."),
        ("NNLS fit", "RMSE 0.1850; R2 0.1799; Pearson r 0.5078", "Partial reconstruction with meaningful but modest global agreement."),
        ("Equal-weight baseline", "RMSE 0.1885; R2 0.1488; Pearson r 0.4165", "NNLS improves RMSE by 1.8% and R2 by 21.0% relative to equal weighting."),
        ("Peaks and residuals", "25 experimental peaks detected; best-match regions 707-779, 870-892, and 921-945 cm-1", "Model captures selected regions but leaves identifiable mismatch bands."),
    ], [2200, 3600, 4760])
    add_p(doc, "The LPS work demonstrates a complete analysis pipeline and produces a chemically interpretable decomposition, but the model should be presented as a preliminary reconstruction rather than a high-fidelity identity proof. The documented mismatch regions motivate adding realistic fragment variants, conformers, linked sugar units, lipid A variants, and possibly learned residual correction.")
    add_figure(doc, ROOT / "20-04" / "lps--main" / "lps--main" / "output" / "plots" / "extended_plot5_nnls_fit_residual_200_2000.png", "Figure 5. LPS NNLS optimized composite spectrum and residual panel. Source: 20-04/lps--main/lps--main/output/plots/extended_plot5_nnls_fit_residual_200_2000.png.")

    add_heading(doc, "5.4 Silver-coated silicon SERS substrate and R6G/NAM testing", 2)
    add_p(doc, "A nanostructured silicon substrate was fabricated and coated with silver through an AgNO3-based deposition route. The later monthly report describes HF etching followed by galvanic displacement deposition of silver nanoparticles. R6G was used as the positive-control molecule before testing NAM, allowing the substrate, Raman acquisition, preprocessing, and calibration workflow to be evaluated independently of the weak biomarker signal.")
    add_table(doc, ["Metric", "Completed result", "Meaning"], [
        ("Substrate", "Nanostructured/etched silicon with Ag coating", "Functional plasmonic test platform established."),
        ("Dataset", "99 spectra in the detailed validation report; approximately 97 spectra in the June progress report", "The folder contains two closely related counts from different reporting versions."),
        ("R6G peak matching", "84.58% accuracy; mean shift error +0.81 cm-1; MAE 3.25 cm-1; RMSE 3.67 cm-1", "Good calibration and positive-control enhancement evidence."),
        ("Preprocessing", "ALS correction, Savitzky-Golay smoothing, cosmic-ray removal, normalization, peak detection, PCA and SNR evaluation", "A complete analysis and QC pipeline was implemented."),
        ("Cosmic-ray processing", "1,491 spikes removed; 99.72% signal retention; 33.57% noise reduction", "Artifact removal was effective without major signal loss according to the report."),
        ("Substrate reproducibility", "Mean Pearson correlation 0.7754; best 0.81; target threshold 0.90", "Good but below the desired reproducibility target."),
        ("R6G optimization", "6 min soaking; 10 power; 5 accumulations selected in the detailed report", "Improved control-molecule SNR and detection frequency."),
        ("NAM response", "NAM features detected but mean SNR 0.89", "NAM is detectable on the current platform but not yet robustly enhanced."),
    ], [2200, 3600, 4760])
    add_callout(doc, "Sensor conclusion", "The substrate and measurement workflow are functional, but the current Ag/Si morphology and/or analyte adsorption are not yet sufficient for reliable NAM biomarker sensing. The folder recommends increasing Ag nanoparticle density, testing multiple deposition cycles, optimizing nanostructure geometry, considering Ag-Au hybrids, and performing concentration-dependent studies.", fill=RED, accent="A33A3A")
    add_source_note(doc, "Primary evidence: Development and Validation of a Silver.docx and Monthly Progress Report-June.docx.")

    add_heading(doc, "5.5 Clinical translation preparation and supporting documentation", 2)
    add_p(doc, "The project archive also contains planning work intended to bridge the laboratory reference studies to clinical validation. These documents are not evidence that clinical samples have already been collected; they are preparatory outputs that define how future work can be executed in a controlled and traceable manner.")
    for item in [
        "NAM stock-solution preparation guidance for initial calibration and low-concentration testing.",
        "SERS experiment phase plans covering stock preparation, deposition, acquisition, controls, and repeatability.",
        "Proposal for a sepsis clinical biobank and SERS diagnostic platform.",
        "Roadmap and SOP for sepsis sample biobanking.",
        "Technical SOP for blood collection, processing, aliquoting, storage, and traceability.",
        "Research questions and risk considerations covering minimum detectable concentration, dilution tolerance, reproducibility, calibration, shelf life, and procurement.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Deliverables Produced", 1)
    add_table(doc, ["Deliverable type", "Examples present in the folder", "Status"], [
        ("Scientific reports", "finalmtp.docx, final-report nam lps lta full details.docx, Project Report_SERS.docx, progress reports", "Multiple report generations completed; final controlled version still to be selected."),
        ("Raman notebooks", "NAM_Raman_Analysis, NAM_Validation_Analysis, NAM_Peak_Assignment, LTA, LPS, and NAM/LTA plotting notebooks", "Analysis workflows documented and outputs generated."),
        ("Python scripts", "cleaning, CSV-to-Excel conversion, plotting, Raman analyzers, Gaussian input generation, validation scripts", "Reusable technical scripts developed across workstreams."),
        ("Processed datasets", "ranked_spectra_table.csv, peak_matching_table.csv, processed spectra, model_metrics.xlsx, peak tables", "Intermediate and final tables exported for analysis and reporting."),
        ("Figures", "Overlays, waterfall plots, heatmaps, residual plots, correlation matrices, annotated spectra", "Publication and presentation-ready figures generated."),
        ("Presentations", "Semester presentations, January/April review decks, NAM/LTA/LPS decks", "Progress communicated through multiple review milestones."),
        ("Protocols and proposals", "NAM stock preparation, SERS experiment phases, biobank proposal, biobank roadmap, technical sample SOP", "Translation and implementation planning documented."),
    ], [1900, 4760, 3900])

    add_heading(doc, "7. Current Status, Limitations, and Evidence Gaps", 1)
    add_heading(doc, "7.1 Workstream status", 2)
    add_table(doc, ["Workstream", "Status", "Evidence-based assessment"], [
        ("Literature, problem definition, and system design", "Completed", "The clinical need, target panel, SERS rationale, workflow, and patent landscape have been documented."),
        ("NAM reference spectrum", "Completed as reference study", "Strong selected-peak agreement; not yet clinical validation or robust global spectral similarity."),
        ("LTA experimental fingerprint", "Completed as reference study", "Good spot reproducibility and characteristic peaks; fragment-model fit requires harmonization and rerun."),
        ("LPS fragment modelling", "Completed preliminary model", "NNLS pipeline and outputs exist; global fit is modest and needs more realistic fragments and replicates."),
        ("Ag/Si SERS substrate", "Functional proof of concept", "R6G validates enhancement and calibration; reproducibility is below target and NAM SNR is low."),
        ("Clinical sample validation", "Not yet completed", "Biobank and SOP planning exist, but the reviewed folder does not establish completed patient-sample validation."),
        ("Multiplex classification / clinical performance", "Not yet completed", "PCA/ML are discussed and some code exists, but no final LOD, sensitivity, specificity, or clinical classifier is established."),
    ], [2500, 1900, 6160])
    add_heading(doc, "7.2 Limitations and quality-control issues", 2)
    for item in [
        "The archive contains duplicate and iterative reports, making it necessary to freeze one authoritative dataset and one final methods description before thesis submission.",
        "Power is represented as percent in NAM/LTA reports and as watts in the enhanced all-data plotting summary. The final report should state the instrument-calibrated unit and preserve the raw metadata.",
        "NAM has excellent peak-position agreement but poor full-spectrum similarity; therefore the result should not be summarized as a perfect overlay or complete spectral match.",
        "LTA has a strong experimental reproducibility branch but also has an unresolved NaN results branch and low whole-spectrum correlation in the expanded report.",
        "LPS NNLS fitting improves on an equal-weight baseline but explains only a limited proportion of global variance; the model is useful for interpretation, not yet as a final diagnostic classifier.",
        "Substrate validation is strongest for R6G, not NAM. The current NAM mean SNR is below the level required for a robust biomarker-sensing claim.",
        "The reviewed evidence is primarily from pure powders, controls, or prepared standards. Real clinical matrices, interference studies, clinical sensitivity/specificity, LOD, batch-to-batch variation, and stability remain open.",
        "Some source documents describe planned machine learning, point-of-care integration, and biobank operations. These should be labelled as proposed work unless supported by executed experiments in the final submission.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Recommended Next Steps", 1)
    steps = [
        ("Freeze and curate the evidence base", "Create a versioned manifest of raw files, processed files, notebooks, figures, and final reports. Choose one canonical dataset for each biomarker and archive superseded branches."),
        ("Standardize the methods", "Resolve the NAM Gaussian/basis-set generations, define power units, fix Raman ranges, record exact normalization and broadening parameters, and make the final methods reproducible from clean paths."),
        ("Optimize the Ag/Si substrate for NAM", "Increase and control Ag coverage, test deposition cycles and Ag-Au alternatives, improve surface morphology and adsorption, and repeat NAM acquisition under a controlled matrix."),
        ("Generate concentration-response evidence", "Measure serial NAM, LTA, and LPS concentrations with blanks, solvent controls, replicates, and matrix-spike recovery. Report LOD, LOQ, calibration curves, linearity, precision, and accuracy."),
        ("Strengthen repeatability and robustness", "Use multiple substrate batches, multiple spatial spots, independent preparation days, and controlled environmental/storage conditions."),
        ("Finalize LTA and LPS models", "Rerun the NaN LTA branch, compare peak-level and global metrics consistently, add conformers/linked fragments where justified, and report residuals with uncertainty."),
        ("Start approved clinical validation", "Use the prepared biobank proposal and SOPs after the required approvals. Build a clinically annotated sample set with healthy/control, bacterial, and relevant non-bacterial cases."),
        ("Develop the multiplex decision layer", "After sufficient reference and clinical data exist, evaluate PCA, SVM, Random Forest, or other models using strict train/test separation and external validation."),
        ("Prepare the final thesis package", "Merge the selected report, methods, figures, tables, and references into one controlled submission document, with a clear distinction between completed evidence and future scope."),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title + ". ")
        set_run_font(r, color=NAVY, bold=True)
        r = p.add_run(body)
        set_run_font(r, color=INK)

    add_heading(doc, "9. Overall Conclusion", 1)
    add_p(doc, "The project has completed the foundational and proof-of-concept stages of a SERS-based sepsis biomarker research programme. It has established a target panel, a modular sensing architecture, controlled Raman acquisition studies, reusable spectral-processing pipelines, DFT reference modelling, fragment-based analysis for complex biomarkers, substrate validation with R6G, and a documented path toward clinical sample testing.")
    add_p(doc, "The project is therefore ready for a focused validation phase rather than a restart. The priority is to convert the existing breadth of work into a controlled, quantitative, and clinically relevant evidence chain: optimized substrate -> concentration response -> reproducibility -> matrix interference -> clinical sample validation -> multiplex classification. Until those steps are completed, the platform should be described as a promising research prototype and reference-analysis framework, not yet as a clinically validated sepsis diagnostic device.")

    add_page_break(doc)
    add_heading(doc, "Appendix A. File Inventory and Evidence Map", 1)
    add_p(doc, "The following map identifies representative evidence files used to build this report. It is intentionally selective for readability; the full folder inventory remains the authoritative archive.")
    add_table(doc, ["Evidence area", "Representative files / folders", "Evidence captured"], [
        ("Project synthesis", "20-04/finalmtp.docx; 20-04/final-report nam lps lta full details.docx; review meeting/finalmtp.docx", "Project objectives, workflow, biomarker chapters, conclusions, novelty, impact, and limitations."),
        ("NAM analysis", "review meeting/NAM-EXP-P/NAM/*.ipynb; output_figures/NAM_analysis_report.txt; output_figures/fig07-fig13*.png", "81 spectra, ranking, best spectrum, 51 peak matches, error statistics, similarity metrics, assignments, and figures."),
        ("LTA analysis", "review meeting/LTA - EXP/Pure Powder spectra/LTA_outputs/*; 20-04/20sec20power-variable location LTA/*; 20-04/lta/*", "136-file ranking, 12-spot reproducibility, DFT fragment strategy, outputs, and unresolved analysis branch."),
        ("LPS analysis", "20-04/lps--main/lps--main/output/final_summary_report.txt; output/plots/*; LPS notebooks", "Five-fragment DFT model, NNLS weights, metrics, peaks, residual regions, and generated tables/figures."),
        ("Substrate validation", "Development and Validation of a Silver.docx; Monthly Progress Report-June.docx", "Ag/Si fabrication, R6G validation, preprocessing, reproducibility, optimization, and NAM limitation."),
        ("Protocols and translation", "Preparation of NAM Stock solution.docx; repot/SERS Experiment phases.docx; review meeting/report doc/*SOP*.docx", "Stock preparation, experiment phases, sample biobank roadmap, and technical sample SOP."),
        ("Molecular modelling support", "molecule building/*; python code/*; review meeting/namfinal laptop simulation/*", "Gaussian/ORCA input generation, molecular structure preparation, Raman extraction, and validation utilities."),
    ], [2000, 4200, 4960])
    add_heading(doc, "Appendix B. Key Quantitative Results", 1)
    add_table(doc, ["Area", "Key values"], [
        ("NAM acquisition and validation", "81 spectra; best 25 s / 80 power; SNR 95.91; 51 matched peaks; peak RMSE 2.115 cm-1; cosine 0.5868; Pearson 0.1710."),
        ("LTA spot reproducibility", "12 spots; 7 characteristic peaks; average CV 16.74%; average correlation 0.9787."),
        ("LTA expanded ranking", "136 spectra; best 20 s / 40 power; SNR 139.03; selected-peak RMSE 2.70 cm-1; global Pearson -0.0084."),
        ("LPS NNLS modelling", "KDO 47.3%; heptose 25.8%; glucosamine 15.3%; myristic acid 11.7%; R2 0.1799; Pearson 0.5078."),
        ("R6G substrate validation", "84.58% peak matching; calibration RMSE 3.67 cm-1; mean correlation 0.7754; 10 power SNR 3.56; 5 accumulations mean SNR 3.70."),
        ("NAM on Ag/Si substrate", "Mean SNR 0.89; spectral features observed but enhancement insufficient for robust sensing."),
    ], [2800, 8360])
    add_heading(doc, "Appendix C. Interpretation Rules Used in This Report", 1)
    for item in [
        "A result is called completed when the folder contains a finished analysis, report, table, or figure with a documented workflow and quantitative output.",
        "A result is called preliminary when the workflow is complete but the model fit, reproducibility, or biological validation is not yet strong enough for a final claim.",
        "A result is called proposed when it appears in a roadmap, SOP, future-work section, or planning document without corresponding executed experimental evidence.",
        "Conflicting values are retained with their source context instead of being averaged or silently reconciled.",
        "Reference powders and standard controls are treated as molecular reference evidence, not as evidence of clinical diagnostic performance.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Submission note", "Before final academic submission, replace this consolidated status report with the institution's required front matter, declaration/certificate pages, approved reference list, and any supervisor-directed formatting. The scientific status and evidence distinctions in this report can be retained as the project-progress chapter.", fill=LIGHT_BLUE, accent=BLUE)

    # Set document metadata.
    props = doc.core_properties
    props.title = "Project Status Report - Sepsis Biomarkers Analysis Using a SERS-Based Sensor"
    props.subject = "Consolidated record of project work completed to date"
    props.author = "Sukesh P"
    props.keywords = "sepsis, SERS, Raman, NAM, LTA, LPS, project status, IIT Jodhpur"
    props.comments = "Prepared from the project folder contents as of 03 August 2026."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
