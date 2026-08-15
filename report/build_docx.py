from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(r"c:\Users\sukes\Downloads\report\Untitled document.md")
OUT = Path(r"c:\Users\sukes\Downloads\report\Final_Project_Report.docx")


def clean_inline(text: str) -> str:
    t = text.strip()
    t = t.replace("**", "")
    t = t.replace("#", "")
    t = t.replace("\\", "")
    t = t.replace("\u200b", "")
    t = t.replace("✅", "")
    t = t.replace("⏳", "")
    t = t.replace("🔮", "")
    t = t.replace("📌", "")
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(re.fullmatch(r"\|?\s*[:\- ]+\|[:\- |]*", s))


def parse_md_table(lines, i):
    if i + 1 >= len(lines):
        return None, i
    if "|" not in lines[i] or not is_table_sep(lines[i + 1]):
        return None, i

    header = [clean_inline(c) for c in lines[i].strip().strip("|").split("|")]
    rows = []
    j = i + 2
    while j < len(lines):
        line = lines[j].rstrip("\n")
        if "|" not in line or not line.strip().startswith("|"):
            break
        row = [clean_inline(c) for c in line.strip().strip("|").split("|")]
        if len(row) < len(header):
            row += [""] * (len(header) - len(row))
        rows.append(row[: len(header)])
        j += 1

    return (header, rows), j - 1


def set_para_run_font(paragraph, size=12, bold=False, italic=False):
    if not paragraph.runs:
        run = paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)
    set_para_run_font(p, size=10)


def set_section_page_number_format(section, fmt: str, start: int = None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def add_section_break(doc):
    p = doc.add_paragraph()
    p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    p._p.pPr.append(sectPr)


def looks_main_section_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s+[A-Z0-9&()\- ,]+$", text))


def looks_subsection(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+(\.\d+)?\s+", text))


def looks_table_caption(text: str) -> bool:
    return bool(re.match(r"^Table\s+\d+\.\d+:", text, flags=re.I))


def looks_figure_caption(text: str) -> bool:
    return bool(re.match(r"^Figure\s+\d+\.\d+:", text, flags=re.I))


def process(source_text: str):
    doc = Document()

    # Page setup: A4, 1-inch margins, portrait
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    lines = source_text.splitlines()

    # Build document body
    i = 0
    in_front_matter = True
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = clean_inline(raw)

        if not line:
            i += 1
            continue

        # table parse from markdown
        parsed, ni = parse_md_table(lines, i)
        if parsed:
            header, rows = parsed
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for c, h in enumerate(header):
                hdr[c].text = h
                for p in hdr[c].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_para_run_font(p, size=12, bold=True)
            for row in rows:
                tr = table.add_row().cells
                for c, v in enumerate(row):
                    tr[c].text = v
                    for p in tr[c].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_para_run_font(p, size=12)
            i = ni + 1
            continue

        # identify transition to main report
        if line.startswith("1. COMPREHENSIVE OVERVIEW") and in_front_matter:
            # section break for Arabic numbering
            doc.add_section(start_type=1)
            main_sec = doc.sections[-1]
            main_sec.page_width = Inches(8.27)
            main_sec.page_height = Inches(11.69)
            main_sec.left_margin = Inches(1)
            main_sec.right_margin = Inches(1)
            main_sec.top_margin = Inches(1)
            main_sec.bottom_margin = Inches(1)
            set_section_page_number_format(main_sec, "decimal", start=1)
            add_page_number_footer(main_sec)
            in_front_matter = False

        p = doc.add_paragraph()

        if line in {"COVER PAGE", "DECLARATION", "CERTIFICATE", "ABSTRACT", "TABLE OF CONTENTS"}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if line == "COVER PAGE" else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            set_para_run_font(p, size=14, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif looks_main_section_heading(line):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(line.upper())
            set_para_run_font(p, size=14, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif looks_subsection(line):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(line)
            set_para_run_font(p, size=12, italic=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif looks_table_caption(line):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)
            set_para_run_font(p, size=12, bold=False)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif looks_figure_caption(line):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)
            set_para_run_font(p, size=12, bold=False)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
        elif re.match(r"^\[\d+\]", line):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(line)
            set_para_run_font(p, size=12)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(line)
            set_para_run_font(p, size=12)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

        i += 1

    # Page numbering setup for front matter
    if len(doc.sections) == 1:
        # fallback: if section split not created
        set_section_page_number_format(doc.sections[0], "lowerRoman", start=1)
        add_page_number_footer(doc.sections[0])
    else:
        front_sec = doc.sections[0]
        set_section_page_number_format(front_sec, "lowerRoman", start=1)
        add_page_number_footer(front_sec)
        # hide page number on cover by separate first page
        front_sec.different_first_page_header_footer = True

    doc.save(OUT)


if __name__ == "__main__":
    src_text = SRC.read_text(encoding="utf-8")
    process(src_text)
    print(str(OUT))
